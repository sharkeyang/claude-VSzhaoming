"""
昭明路线图.md → DOCX 转换脚本
用法: python md2docx.py [输入文件.md]
  默认: _规则文档/全局_C昭明路线图.md
"""

import pathlib, re, datetime, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import markdown
from bs4 import BeautifulSoup

# 输入文件
BASE = pathlib.Path(__file__).parent.parent
默认输入 = BASE / '_规则文档' / '全局_C昭明路线图.md'
输入文件 = pathlib.Path(sys.argv[1]) if len(sys.argv) > 1 else 默认输入
输出目录 = 输入文件.parent
ts = datetime.datetime.now().strftime('%Y%m%d_%H%M')

# 读 .md
md = 输入文件.read_text(encoding='utf-8')

# 移除 <style> 和 YAML 前注
md = re.sub(r'^<style>.*?</style>\n?', '', md, flags=re.DOTALL)
md = re.sub(r'^---.*?---\n?', '', md, flags=re.DOTALL)
md = md.strip()

# md→html
html = markdown.markdown(md, extensions=['tables', 'fenced_code', 'codehilite'])
soup = BeautifulSoup(html, 'html.parser')

# 建 Word
doc = Document()

def 设中文字体(样式, 字体名='SimSun', 大小=Pt(10), 加粗=False):
    style = doc.styles[样式]
    style.font.name = 字体名
    style.font.size = 大小
    style.font.bold = 加粗
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {
        qn('w:eastAsia'): 字体名, qn('w:ascii'): 字体名, qn('w:hAnsi'): 字体名
    })
    rPr.insert(0, rFonts)

def 设段字体(段落, 字体名='SimSun', 大小=Pt(10)):
    for run in 段落.runs:
        run.font.name = 字体名
        run.font.size = 大小

# 默认字体
设中文字体('Normal', 'SimSun', Pt(10))
for i in range(1, 5):
    设中文字体(f'Heading {i}', 'SimHei', {1: Pt(15), 2: Pt(13), 3: Pt(12), 4: Pt(11)}[i], True)

# 逐元素处理
for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'p', 'pre', 'table', 'ul', 'ol', 'hr', 'blockquote']):
    try:
        if tag.name in 'h12345':
            doc.add_heading(tag.get_text().strip(), level=int(tag.name[1]))
        elif tag.name == 'p':
            t = tag.get_text().strip()
            if t:
                doc.add_paragraph(t)
        elif tag.name == 'pre':
            p = doc.add_paragraph()
            run = p.add_run(tag.get_text())
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
        elif tag.name == 'table':
            rows = tag.find_all('tr')
            if rows:
                ncols = max(len(r.find_all(['th','td'])) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=ncols)
                tbl.style = 'Table Grid'
                for i, row in enumerate(rows):
                    cells = row.find_all(['th','td'])
                    for j, cell in enumerate(cells):
                        if j < ncols:
                            run = tbl.rows[i].cells[j].paragraphs[0].add_run(
                                cell.get_text().strip())
                            run.font.size = Pt(8)
                            run.font.name = 'SimSun'
        elif tag.name in ['ul', 'ol']:
            for li in tag.find_all('li'):
                t = li.get_text().strip()
                if t:
                    p = doc.add_paragraph(t, style='List Bullet')
                    设段字体(p, 'SimSun', Pt(10))
        elif tag.name == 'hr':
            p = doc.add_paragraph('─' * 40)
        elif tag.name == 'blockquote':
            for pt in tag.find_all('p'):
                t = pt.get_text().strip()
                if t:
                    p = doc.add_paragraph()
                    run = p.add_run(f'  {t}')
                    run.font.size = Pt(9)
                    run.font.italic = True
    except:
        pass

输出路径 = 输出目录 / f'{输入文件.stem}_{ts}.docx'
doc.save(str(输出路径))
print(f'✅ {输出路径.name} ({(输出路径.stat().st_size / 1024):.0f}KB)')